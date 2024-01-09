import re
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import _Cell

def open_document(username):
    try:
        document = Document(f"C:/Users/{username}/Desktop/Cautela Template.docx")
        return document
    except:
        print("O arquivo \"Cautela Template.docx\" deve estar na área de trabalho")
        exit()

def create_item_table_paragraph_style(document):
    styles = document.styles
    item_table_style = styles.add_style('Table Item', WD_STYLE_TYPE.PARAGRAPH)
    font = item_table_style.font
    font.name = 'Calibri'
    font.size = Pt(11)

def docx_replace_regex(doc_obj, regex , replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)

def replace_words_from_dictionary(document, dictionary):
     for word, replacement in dictionary.items():
        word_re = re.compile(word)
        docx_replace_regex(document, word_re, replacement)

def add_item_to_item_table(document, item_information):

    item_table = document.tables[1]
    row = item_table.add_row()
    j = 0

    for cell in row.cells:
        set_cell_border(
            cell, 
            top = {"val": "single", "color": "#000000"},
            bottom = {"val": "single", "color": "#000000"},
            start = {"val": "single", "color": "#000000"},
            end = {"val": "single", "color": "#000000"},
        )
        paragraph = cell.paragraphs[0]
        paragraph.text = item_information[j]
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.style = 'Table Item'
        j += 1

def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))